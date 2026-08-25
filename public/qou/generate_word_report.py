import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from generate_super_dashboard import extract_comprehensive_data

def set_rtl(element):
    """Set the Right-To-Left (RTL) property on a paragraph."""
    pPr = element._element.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)

def set_table_rtl(table):
    """Set the RTL property for a table."""
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        bidiVisual = OxmlElement('w:bidiVisual')
        tblPr[0].append(bidiVisual)

def add_rtl_paragraph(doc, text, style=None, bold=False, size=None, align=WD_ALIGN_PARAGRAPH.RIGHT):
    p = doc.add_paragraph(style=style)
    p.alignment = align
    set_rtl(p)
    run = p.add_run(text)
    if bold:
        run.bold = True
    if size:
        run.font.size = Pt(size)
    run.font.rtl = True
    return p

def generate_word_report():
    print("استخراج البيانات لإعداد تقرير الوورد...")
    employees, branches_stats = extract_comprehensive_data()
    
    # Sort branches by total minutes
    branches_stats = sorted(branches_stats, key=lambda x: x['إجمالي الدقائق'], reverse=True)
    
    doc = Document()
    
    # Add Title
    title = add_rtl_paragraph(doc, 'التقرير الشامل لمتابعة تأخير الموظفين', bold=True, size=24, align=WD_ALIGN_PARAGRAPH.CENTER)
    
    # Overall summary
    add_rtl_paragraph(doc, 'ملخص عام للفروع:', bold=True, size=16)
    
    total_emps = len(employees)
    total_mins = sum(b['إجمالي الدقائق'] for b in branches_stats)
    total_delays = sum(b['إجمالي حالات التأخير'] for b in branches_stats)
    
    add_rtl_paragraph(doc, f'إجمالي الموظفين المتأخرين: {total_emps}')
    add_rtl_paragraph(doc, f'إجمالي دقائق التأخير: {total_mins}')
    add_rtl_paragraph(doc, f'إجمالي الساعات: {round(total_mins/60, 2)}')
    add_rtl_paragraph(doc, f'إجمالي حالات التأخير: {total_delays}')
    
    doc.add_page_break()
    
    # Branch Summary Table
    add_rtl_paragraph(doc, 'ملخص أداء الفروع', bold=True, size=18, align=WD_ALIGN_PARAGRAPH.CENTER)
    
    table = doc.add_table(rows=1, cols=8)
    table.style = 'Table Grid'
    set_table_rtl(table)
    
    headers = ['الفرع', 'إجمالي الدقائق', 'إجمالي الساعات', 'متوسط الدقائق', 'الموظفين المتأخرين', 'الموظف الأكثر تأخيراً', 'إجمالي الحالات', 'اليوم الأكثر تكراراً']
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        p = hdr_cells[i].paragraphs[0]
        p.text = h
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_rtl(p)
        for run in p.runs:
            run.font.bold = True
            run.font.rtl = True
            
    for b in branches_stats:
        row_cells = table.add_row().cells
        data = [
            b['الفرع'], str(b['إجمالي الدقائق']), str(b['إجمالي الساعات']), 
            str(b['متوسط دقائق التأخير']), str(b['إجمالي الموظفين المتأخرين']), 
            b['الموظف الأكثر تأخيراً'], str(b['إجمالي حالات التأخير']), 
            b['اليوم الأكثر تكراراً']
        ]
        for i, val in enumerate(data):
            p = row_cells[i].paragraphs[0]
            p.text = val
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_rtl(p)
            for run in p.runs:
                run.font.rtl = True
                
    doc.add_page_break()
    
    # Detailed Branch Reports
    for b in branches_stats:
        branch_name = b['الفرع']
        add_rtl_paragraph(doc, f'تقرير تفصيلي لفرع: {branch_name}', bold=True, size=18, align=WD_ALIGN_PARAGRAPH.CENTER)
        
        # Branch specific summary
        add_rtl_paragraph(doc, f"عدد الموظفين المتأخرين: {b['إجمالي الموظفين المتأخرين']} | إجمالي الدقائق: {b['إجمالي الدقائق']} | إجمالي الحالات: {b['إجمالي حالات التأخير']}", bold=True, size=12)
        
        # Branch employees
        branch_emps = sorted([e for e in employees if e['الفرع'] == branch_name], key=lambda x: x['إجمالي الدقائق'], reverse=True)
        
        emp_table = doc.add_table(rows=1, cols=7)
        emp_table.style = 'Light Shading Accent 1'
        set_table_rtl(emp_table)
        
        emp_headers = ['م', 'اسم الموظف', 'الدقائق', 'الحالات', 'الساعات', 'الأشهر المسجلة', 'اليوم الأكثر تكراراً']
        hdr_cells = emp_table.rows[0].cells
        for i, h in enumerate(emp_headers):
            p = hdr_cells[i].paragraphs[0]
            p.text = h
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_rtl(p)
            for run in p.runs:
                run.font.bold = True
                run.font.rtl = True
                
        for idx, emp in enumerate(branch_emps):
            row_cells = emp_table.add_row().cells
            data = [
                str(idx+1), emp['اسم الموظف'], str(emp['إجمالي الدقائق']),
                str(emp['إجمالي حالات التأخير']), str(emp['إجمالي الساعات']),
                emp['الأشهر المسجلة'], emp['اليوم الأكثر تكراراً']
            ]
            for i, val in enumerate(data):
                p = row_cells[i].paragraphs[0]
                p.text = val
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_rtl(p)
                for run in p.runs:
                    run.font.rtl = True
        
        doc.add_paragraph() # Spacing
    
    output_path = os.path.abspath('تقرير_المتابعة_الشامل_v3.docx')
    doc.save(output_path)
    print("تم إنشاء ملف الوورد بنجاح!")

if __name__ == "__main__":
    generate_word_report()
