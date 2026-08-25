import os
import glob
import re
import unicodedata
from collections import Counter
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pdfplumber

def parse_pdf_comprehensive(filepath):
    branch_name = os.path.basename(filepath).replace('.pdf', '')
    employees_dict = {}
    current_employee_name = None
    days_counter = Counter()
    total_delays = 0
    total_minutes = 0

    with pdfplumber.open(filepath) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                lines = text.split('\n')
                for line in lines:
                    rev_line = unicodedata.normalize('NFKC', line[::-1])
                    
                    if "الموظف:" in rev_line and "عدد التأخيرات:" in rev_line:
                        match = re.search(r"الموظف:\s*(.*?)\s*عدد التأخيرات:", rev_line)
                        if match:
                            name = match.group(1).strip()
                            if name in ["ماهر عبد الحميد يونس أبو ساكور", "محمد سعدي عبد الرحمن ابو ساكور", "محمد سعدي عبد الرحمن أبو ساكور"]:
                                current_employee_name = None
                                continue
                            current_employee_name = name
                            if name not in employees_dict:
                                employees_dict[name] = {
                                    'اسم الموظف': name,
                                    'عدد حالات التأخير': 0,
                                    'إجمالي دقائق التأخير': 0,
                                    'الأشهر': set()
                                }
                            
                            delay_match = re.search(r"عدد التأخيرات:\s*(\d+)", rev_line)
                            if delay_match:
                                employees_dict[name]['عدد حالات التأخير'] = int(delay_match.group(1)[::-1])
                    
                    if "مجموع دقائق التأخير" in rev_line:
                        match = re.search(r"مجموع دقائق التأخير\s*(\d+)", rev_line)
                        if match and current_employee_name is not None:
                            employees_dict[current_employee_name]['إجمالي دقائق التأخير'] = int(match.group(1)[::-1])

                    # Try finding date in the normal or reversed line to extract month
                    # format DD/MM/YYYY or YYYY/MM/DD
                    # normal line:
                    for d in re.finditer(r'(\d{2})/(\d{2})/(\d{4})', line):
                        # d.group(2) is month
                        if current_employee_name is not None:
                            employees_dict[current_employee_name]['الأشهر'].add(d.group(2))
                    for d in re.finditer(r'(\d{4})/(\d{2})/(\d{2})', line):
                        # d.group(2) is month
                        if current_employee_name is not None:
                            employees_dict[current_employee_name]['الأشهر'].add(d.group(2))
                            
            tables = page.extract_tables()
            for table in tables:
                for row in table:
                    if row and len(row) >= 4 and row[0] and str(row[0]).isdigit():
                        total_delays += 1
                        day_str = row[1]
                        if day_str:
                            day_clean = unicodedata.normalize('NFKC', day_str[::-1]).strip()
                            days_counter[day_clean] += 1
                        
                        # Fallback for dates in table if not caught by text
                        date_str = str(row[3]) if len(row) > 3 else ""
                        if date_str and current_employee_name is not None:
                            m = re.search(r'(\d{2})/(\d{2})/(\d{4})', date_str)
                            if m:
                                # We can't guarantee current_employee_name is perfectly matched in table loop if multiple emps per page, 
                                # but it's a good fallback
                                employees_dict[current_employee_name]['الأشهر'].add(m.group(2))

    calc_total_delays = sum(e['عدد حالات التأخير'] for e in employees_dict.values())
    calc_total_minutes = sum(e['إجمالي دقائق التأخير'] for e in employees_dict.values())
    
    # Calculate hours and join months
    for e in employees_dict.values():
        e['إجمالي الساعات'] = round(e['إجمالي دقائق التأخير'] / 60, 2)
        e['الأشهر المسجلة'] = "، ".join(sorted(list(e['الأشهر']))) if e['الأشهر'] else "غير محدد"

    top_emp = max(employees_dict.values(), key=lambda x: x['إجمالي دقائق التأخير']) if employees_dict else None
    top_day = days_counter.most_common(1)[0] if days_counter else ("لا يوجد", 0)

    stats = {
        'الفرع': branch_name,
        'إجمالي الموظفين المتأخرين': len(employees_dict),
        'إجمالي حالات التأخير': calc_total_delays,
        'إجمالي الدقائق': calc_total_minutes,
        'إجمالي الساعات': round(calc_total_minutes / 60, 2),
        'متوسط دقائق التأخير': round(calc_total_minutes / calc_total_delays if calc_total_delays > 0 else 0, 2),
        'الموظف الأكثر تأخيراً': f"{top_emp['اسم الموظف']} ({top_emp['إجمالي الساعات']} ساعة)" if top_emp else "لا يوجد",
        'اليوم الأكثر تكراراً': f"{top_day[0]} ({top_day[1]} مرة)"
    }
    
    return branch_name, stats, list(employees_dict.values())

def generate_all_reports():
    files = glob.glob("*.pdf")
    all_branches_stats = []
    all_employees_data = {}
    
    overall_total_employees = 0
    overall_total_delays = 0
    overall_total_minutes = 0
    overall_days_counter = Counter()
    overall_emp_minutes = {}

    for f in files:
        branch, stats, employees = parse_pdf_comprehensive(f)
        all_branches_stats.append(stats)
        all_employees_data[branch] = employees
        
        overall_total_employees += stats['إجمالي الموظفين المتأخرين']
        overall_total_delays += stats['إجمالي حالات التأخير']
        overall_total_minutes += stats['إجمالي الدقائق']
        
        for e in employees:
            overall_emp_minutes[f"{e['اسم الموظف']} ({branch})"] = e['إجمالي دقائق التأخير']
            
        with pdfplumber.open(f) as pdf:
            for page in pdf.pages:
                for table in page.extract_tables():
                    for row in table:
                        if row and len(row)>=4 and row[0] and str(row[0]).isdigit() and row[1]:
                            day_clean = unicodedata.normalize('NFKC', row[1][::-1]).strip()
                            overall_days_counter[day_clean] += 1

    overall_top_emp = max(overall_emp_minutes.items(), key=lambda x: x[1]) if overall_emp_minutes else ("لا يوجد", 0)
    overall_top_day = overall_days_counter.most_common(1)[0] if overall_days_counter else ("لا يوجد", 0)
    overall_avg = round(overall_total_minutes / overall_total_delays if overall_total_delays > 0 else 0, 2)
    overall_total_hours = round(overall_total_minutes / 60, 2)

    # ================= 1. EXCEL DASHBOARD =================
    writer = pd.ExcelWriter('التقرير_النهائي_الشامل.xlsx', engine='xlsxwriter')
    workbook = writer.book

    title_format = workbook.add_format({'bold': True, 'font_size': 20, 'font_color': '#FFFFFF', 'bg_color': '#1F4E78', 'align': 'center', 'valign': 'vcenter', 'border': 1})
    header_format = workbook.add_format({'bold': True, 'font_color': '#FFFFFF', 'bg_color': '#2F75B5', 'align': 'center', 'valign': 'vcenter', 'border': 1})
    cell_format = workbook.add_format({'align': 'center', 'valign': 'vcenter', 'border': 1})
    stat_box_title = workbook.add_format({'bold': True, 'font_size': 12, 'font_color': '#FFFFFF', 'bg_color': '#C00000', 'align': 'center', 'valign': 'vcenter', 'border': 1})
    stat_box_value = workbook.add_format({'bold': True, 'font_size': 14, 'bg_color': '#F2F2F2', 'align': 'center', 'valign': 'vcenter', 'border': 1})

    worksheet = workbook.add_worksheet('الواجهة الرئيسية (الملخص)')
    worksheet.right_to_left()
    
    worksheet.merge_range('B2:I3', 'لوحة القيادة الرئيسية - إحصائيات التأخيرات الصباحية', title_format)

    worksheet.write('B5', 'إجمالي الموظفين', stat_box_title)
    worksheet.write('B6', overall_total_employees, stat_box_value)
    worksheet.write('C5', 'إجمالي التأخيرات', stat_box_title)
    worksheet.write('C6', overall_total_delays, stat_box_value)
    worksheet.write('D5', 'إجمالي الدقائق', stat_box_title)
    worksheet.write('D6', f"{overall_total_minutes:,}", stat_box_value)
    worksheet.write('E5', 'إجمالي الساعات', stat_box_title)
    worksheet.write('E6', overall_total_hours, stat_box_value)
    worksheet.write('F5', 'متوسط التأخير (دقيقة)', stat_box_title)
    worksheet.write('F6', overall_avg, stat_box_value)
    worksheet.merge_range('G5:H5', 'الموظف الأكثر تأخيراً', stat_box_title)
    worksheet.merge_range('G6:H6', f"{overall_top_emp[0]} ({round(overall_top_emp[1]/60, 2)} ساعة)", stat_box_value)
    worksheet.write('I5', 'اليوم الأكثر تكراراً', stat_box_title)
    worksheet.write('I6', f"{overall_top_day[0]} ({overall_top_day[1]})", stat_box_value)

    worksheet.merge_range('B9:I9', 'إحصائيات تفصيلية لكل فرع / دائرة', title_format)
    
    df_branches = pd.DataFrame(all_branches_stats)
    df_branches = df_branches.sort_values('إجمالي الدقائق', ascending=False)
    df_branches.insert(0, 'الترتيب', range(1, len(df_branches) + 1))
    
    for col_num, col_name in enumerate(df_branches.columns):
        worksheet.write(10, col_num + 1, col_name, header_format)
    for row_num, row_data in enumerate(df_branches.values):
        for col_num, cell_data in enumerate(row_data):
            worksheet.write(row_num + 11, col_num + 1, cell_data, cell_format)

    worksheet.set_column('B:B', 10)
    worksheet.set_column('C:C', 35)
    worksheet.set_column('D:J', 20)

    for branch, employees in all_employees_data.items():
        if not employees: continue
        sheet_name = branch[:31]
        ws = workbook.add_worksheet(sheet_name)
        ws.right_to_left()
        ws.merge_range('A1:G2', f'الإحصائيات التفصيلية للموظفين - {branch}', title_format)
        
        # Prepare DataFrame
        clean_emps = []
        for e in employees:
            clean_emps.append({
                'اسم الموظف': e['اسم الموظف'],
                'إجمالي الساعات': e['إجمالي الساعات'],
                'إجمالي دقائق التأخير': e['إجمالي دقائق التأخير'],
                'عدد حالات التأخير': e['عدد حالات التأخير'],
                'متوسط دقائق التأخير للحالة': round(e['إجمالي دقائق التأخير'] / e['عدد حالات التأخير'], 2) if e['عدد حالات التأخير'] > 0 else 0,
                'الأشهر المسجلة': e['الأشهر المسجلة']
            })
            
        df = pd.DataFrame(clean_emps)
        df = df.sort_values('إجمالي الساعات', ascending=False)
        df.insert(0, 'الترتيب', range(1, len(df) + 1))
        
        for col_num, col_name in enumerate(df.columns):
            ws.write(3, col_num, col_name, header_format)
        for row_num, row_data in enumerate(df.values):
            for col_num, cell_data in enumerate(row_data):
                ws.write(row_num + 4, col_num, cell_data, cell_format)
                
        ws.set_column('B:B', 35)
        ws.set_column('C:G', 20)

    writer.close()
    
    # ================= 2. WORD REPORT =================
    doc = Document()
    title = doc.add_heading('التقرير الإداري المفصل لتأخيرات الموظفين', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    all_stats_sorted = sorted(all_branches_stats, key=lambda x: x['إجمالي الدقائق'], reverse=True)
    
    doc.add_heading('الملخص العام', level=1)
    p = doc.add_paragraph()
    p.add_run(f"إجمالي عدد الموظفين المتأخرين: {overall_total_employees}\n").bold = True
    p.add_run(f"إجمالي عدد التأخيرات: {overall_total_delays}\n").bold = True
    p.add_run(f"إجمالي الساعات المهدرة: {overall_total_hours} ساعة ").bold = True
    p.add_run(f"({overall_total_minutes:,} دقيقة)\n")
    if overall_total_delays > 0:
        p.add_run(f"متوسط التأخير للحالة الواحدة: {overall_avg} دقيقة\n").bold = True
        
    doc.add_page_break()
    
    for stats in all_stats_sorted:
        branch = stats['الفرع']
        doc.add_heading(f"فرع / دائرة: {branch}", level=2)
        p2 = doc.add_paragraph()
        p2.add_run(f"إجمالي الموظفين: {stats['إجمالي الموظفين المتأخرين']} | إجمالي الساعات: {stats['إجمالي الساعات']} | إجمالي التأخيرات: {stats['إجمالي حالات التأخير']}")
        
        emps = all_employees_data[branch]
        if emps:
            # Sort emps for word doc
            emps = sorted(emps, key=lambda x: x['إجمالي الساعات'], reverse=True)
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'اسم الموظف'
            hdr_cells[1].text = 'إجمالي الساعات'
            hdr_cells[2].text = 'إجمالي الدقائق'
            hdr_cells[3].text = 'التأخيرات'
            hdr_cells[4].text = 'الأشهر'
            
            for emp in emps[:10]: 
                row_cells = table.add_row().cells
                row_cells[0].text = emp['اسم الموظف']
                row_cells[1].text = str(emp['إجمالي الساعات'])
                row_cells[2].text = str(emp['إجمالي دقائق التأخير'])
                row_cells[3].text = str(emp['عدد حالات التأخير'])
                row_cells[4].text = emp['الأشهر المسجلة']
                
            if len(emps) > 10:
                doc.add_paragraph(f"... وموظفين آخرين ({len(emps)-10} موظف).")
        doc.add_paragraph()
        
    doc.save('التقرير_الإداري_للتأخيرات.docx')
    print("All reports generated successfully.")

if __name__ == "__main__":
    generate_all_reports()
